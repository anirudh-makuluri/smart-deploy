"""Generate docs/SMARTDEPLOY_OVERVIEW.docx from the overview content and screenshots."""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt

ROOT = Path(__file__).resolve().parents[1]
ASSETS = ROOT / "docs" / "assets" / "smartdeploy"
OUTPUT = ROOT / "docs" / "SMARTDEPLOY_OVERVIEW.docx"


def add_heading(doc: Document, text: str, level: int) -> None:
    doc.add_heading(text, level=level)


def add_bullet_list(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def add_image(doc: Document, filename: str, caption: str, width_inches: float = 6.25) -> None:
    image_path = ASSETS / filename
    if not image_path.exists():
        doc.add_paragraph(f"[Missing image: {filename}]")
        return
    doc.add_picture(str(image_path), width=Inches(width_inches))
    caption_paragraph = doc.add_paragraph(caption)
    caption_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if caption_paragraph.runs:
        caption_paragraph.runs[0].italic = True
        caption_paragraph.runs[0].font.size = Pt(9)


def build_document() -> Document:
    doc = Document()

    title = doc.add_heading("SmartDeploy", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT

    subtitle = doc.add_paragraph("AI-Native AWS Deployment Platform")
    subtitle.runs[0].bold = True

    stack_line = doc.add_paragraph(
        "Currently built on ECS Fargate, Lambda, SQS, ECR, ALB, CloudWatch, Bedrock, and AWS Systems Manager."
    )
    stack_line.runs[0].italic = True

    meta = doc.add_paragraph()
    meta.add_run("Built by: ").bold = True
    meta.add_run("Makuluri Labs\n")
    meta.add_run("Founder: ").bold = True
    meta.add_run("Anirudh Makuluri\n")
    meta.add_run("Website: ").bold = True
    meta.add_run("https://makuluri.com")

    add_heading(doc, "Overview", 1)
    doc.add_paragraph(
        "SmartDeploy is an AI-native AWS deployment platform that analyzes GitHub repositories, "
        "generates deployment blueprints, automates production deployments, and includes an AI "
        "deployment agent and CLI for troubleshooting and operations."
    )
    doc.add_paragraph(
        "Developers connect a GitHub repository, and SmartDeploy handles the full path from "
        "codebase analysis to live AWS infrastructure—without requiring manual cloud configuration."
    )

    add_heading(doc, "Why AWS?", 1)
    doc.add_paragraph("AWS is the core platform behind SmartDeploy.")
    doc.add_paragraph("The platform currently leverages:")
    add_bullet_list(
        doc,
        [
            "Amazon ECS Fargate",
            "AWS Lambda",
            "Amazon SQS",
            "Amazon ECR",
            "Application Load Balancer",
            "Amazon S3",
            "Amazon CloudWatch",
            "AWS Systems Manager",
            "AWS IAM",
            "Amazon Bedrock",
        ],
    )
    doc.add_paragraph(
        "SmartDeploy uses Amazon SQS and AWS Lambda to decouple deployment execution from the "
        "control plane, allowing long-running deployments to execute asynchronously while keeping "
        "the application responsive."
    )
    doc.add_paragraph(
        "Unlike a traditional SaaS application, SmartDeploy also provisions AWS infrastructure "
        "for its users. Every deployment executed through SmartDeploy creates and manages AWS "
        "resources on behalf of developers."
    )
    doc.add_paragraph("AWS Activate credits will allow us to:")
    add_bullet_list(
        doc,
        [
            "Scale deployment capacity",
            "Expand the AI Deployment Agent",
            "Expand the SmartDeploy CLI",
            "Perform larger-scale deployment and reliability testing",
            "Continue building production-grade AWS infrastructure while scaling SmartDeploy",
        ],
    )

    add_heading(doc, "Product Highlights", 1)

    add_heading(doc, "Repository Dashboard", 2)
    doc.add_paragraph(
        "SmartDeploy organizes repositories into deployment workspaces, automatically detects "
        "deployable services, and tracks deployment history across projects."
    )
    add_image(
        doc,
        "repository-dashboard.png",
        "SmartDeploy Repository Dashboard — deployments workspace with service detection across connected GitHub repositories",
    )

    add_heading(doc, "AI Build Analysis", 2)
    doc.add_paragraph(
        "SmartDeploy analyzes GitHub repositories, validates build plans, detects deployment "
        "issues before deployment, and allows developers to modify deployment commands before execution."
    )
    add_image(
        doc,
        "build-analysis.png",
        "SmartDeploy Build Analysis — scan results with build validation, Railpack detection, and editable install/build commands",
    )

    add_heading(doc, "Deployment Blueprint", 2)
    doc.add_paragraph(
        "Before deployment, SmartDeploy generates a visual deployment blueprint showing every "
        "deployment stage, required AWS resources, networking configuration, runtime settings, "
        "and deployment flow."
    )
    doc.add_paragraph(
        "This provides complete visibility into what will be created before any AWS resources are provisioned."
    )
    add_image(
        doc,
        "deployment-blueprint.png",
        "SmartDeploy Deployment Blueprint — preview of AUTH, BUILD, SETUP, DEPLOY, and DONE stages from GitHub to ECS Fargate",
    )

    add_heading(doc, "Current Focus", 1)
    doc.add_paragraph("SmartDeploy is actively expanding in the following areas:")
    add_bullet_list(
        doc,
        [
            "AI Deployment Agent",
            "SmartDeploy CLI",
            "Deployment Observability",
            "Automated Deployment Remediation",
            "Deployment Orchestration",
            "Scalable Deployment Infrastructure",
        ],
    )

    add_heading(doc, "Current Status", 1)
    add_bullet_list(
        doc,
        [
            "Bootstrapped and founder-funded",
            "Public product under active development",
            "AI-native AWS deployment platform",
            "Deploys customer applications directly onto AWS infrastructure",
            "Designed to simplify production deployments while providing complete deployment transparency",
        ],
    )

    add_heading(doc, "Why AWS Activate?", 1)
    doc.add_paragraph(
        "SmartDeploy has been entirely self-funded to date, with AWS infrastructure costs covered "
        "personally during development. AWS Activate credits would accelerate product development by "
        "enabling larger-scale deployment testing, expanding the AI deployment agent and CLI, and "
        "continuing to build production-grade AWS infrastructure without constraining experimentation."
    )
    doc.add_paragraph(
        "The long-term goal is to make SmartDeploy the simplest and most transparent way for developers "
        "to deploy production applications on AWS while providing an intelligent deployment agent that "
        "assists with deployment planning, troubleshooting, and operational workflows."
    )

    return doc


def main() -> None:
    document = build_document()
    document.save(OUTPUT)
    print(f"Wrote {OUTPUT}")


if __name__ == "__main__":
    main()
